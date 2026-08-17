"""Сборка .docx из .md для патентного пакета — замена pandoc.

Зачем: `PROMPT_собрать_docx.md` предписывает pandoc, но на машине его нет.
Здесь тот же результат на python-docx.

⚠️ Интерпретатор: нужен тот, где стоит python-docx — системный Python 3.12,
НЕ .venv проекта (там пакета нет):

    "F:/Program Files (x86)/Python312/python.exe" _md2docx.py

⚠️ ГРАБЛЯ из `PROMPT_собрать_docx.md`: bash-mount отстаёт от записи на хост,
и свежий .md через bash читается обрезанным. Поэтому скрипт запускается
ВИНДОВЫМ python (читает хостовую ФС напрямую) и перед конвертацией печатает
число строк и хвост каждого файла — сверять глазами.

Канонические .md и .docx предыдущей редакции лежат уровнем выше и НЕ трогаются —
эта папка самодостаточна: и исходники, и собранные документы новой редакции.

Поддержано: заголовки, абзацы, **жирный**/`код` inline, таблицы, цитаты,
списки, блоки кода, картинки, горизонтальные линии.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

PATENT = Path(__file__).resolve().parent   # документы лежат рядом со скриптом

# исходник → результат (обе стороны в этой же папке)
JOBS = [
    ("zayavka_SPOSOB.md", "Заявка_СПОСОБ_3FFT_ред_2026-08-09.docx"),
    ("zayavka_USTROJSTVO.md", "Заявка_УСТРОЙСТВО_3FFT_ред_2026-08-09.docx"),
    ("formula_izobreteniya.md", "Формула_3FFT_ред_2026-08-09.docx"),
    ("bloki_ustroystva_20.md", "Блоки_устройства_ред_2026-08-09.docx"),
    
    ("Дополнение_кластер_2026-08-17.md", "Дополнение_кластер_2026-08-17.docx"),
    ("Письмо_поверенному_2026-08-09.md", "Письмо_поверенному_2026-08-09.docx"),
    ("Сопроводительное_письмо_2026-08-09.md", "Сопроводительное_письмо_2026-08-09.docx"),
]

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")


def add_runs(par, text: str) -> None:
    """Разбирает **жирный**, `код`, *курсив* внутри абзаца."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            par.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = par.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            par.add_run(part[1:-1]).italic = True
        else:
            par.add_run(part)


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").split("\n")
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()

        # блок кода
        if s.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(buf))
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
            continue

        # таблица
        if s.startswith("|") and i + 1 < n and re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):
            head = split_row(s)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            t = doc.add_table(rows=1, cols=len(head))
            t.style = "Table Grid"
            for c, txt in zip(t.rows[0].cells, head):
                c.text = ""
                add_runs(c.paragraphs[0], txt)
                for run in c.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = t.add_row().cells
                for c, txt in zip(cells, row[: len(head)]):
                    c.text = ""
                    add_runs(c.paragraphs[0], txt)
            doc.add_paragraph()
            continue

        # картинка
        m = re.match(r"^!\[(.*?)\]\((.+?)\)$", s)
        if m:
            img = (md_path.parent / m.group(2)).resolve()
            if img.exists():
                try:
                    doc.add_picture(str(img), width=Inches(6.0))
                    cap = doc.add_paragraph(m.group(1))
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(10)
                except Exception as e:                       # noqa: BLE001
                    doc.add_paragraph(f"[не вставлена картинка {m.group(2)}: {e}]")
            else:
                p = doc.add_paragraph()
                r = p.add_run(f"[ФАЙЛ РИСУНКА ОТСУТСТВУЕТ: {m.group(2)} — {m.group(1)}]")
                r.bold = True
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            i += 1
            continue

        # заголовок
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            doc.add_heading(re.sub(r"[*`]", "", m.group(2)), level=min(len(m.group(1)), 4))
            i += 1
            continue

        # горизонтальная линия / YAML-обвязка
        if s in ("---", "***", "___"):
            i += 1
            continue
        if s.startswith("title:") and s.endswith('"'):
            i += 1
            continue

        # цитата
        if s.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            add_runs(p, " ".join(x for x in buf if x))
            for r in p.runs:
                r.italic = True
                r.font.size = Pt(11)
            continue

        # список
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            style = "List Bullet" if m.group(2) in ("-", "*") else "List Number"
            p = doc.add_paragraph(style=style)
            add_runs(p, m.group(3))
            i += 1
            continue

        # пустая
        if not s:
            i += 1
            continue

        # обычный абзац (склеиваем до пустой строки)
        buf = []
        while i < n and lines[i].strip() and not re.match(
            r"^\s*(#{1,6}\s|\||>|```|!\[|[-*]\s|\d+\.\s|---$)", lines[i]
        ):
            buf.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, " ".join(buf))

    doc.save(str(docx_path))


def main() -> int:
    print(f"Рабочая папка: {PATENT}\n")
    ok = True
    for src, dst in JOBS:
        sp, dp = PATENT / src, PATENT / dst
        if not sp.exists():
            print(f"[ПРОПУСК] нет исходника: {src}")
            ok = False
            continue

        text = sp.read_text(encoding="utf-8")
        print(f"[{src}]")
        print(f"   строк: {text.count(chr(10))} | символов: {len(text)}")
        print(f"   хвост: {text.strip()[-70:]!r}")

        convert(sp, dp)
        print(f"   → {dst}  ({dp.stat().st_size // 1024} КБ)\n")
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
